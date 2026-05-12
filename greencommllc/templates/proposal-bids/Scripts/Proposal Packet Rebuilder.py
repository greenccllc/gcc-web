"""
Proposal Packet Rebuilder — GCC LV Div27-28

Batch-applies GCC brand formatting to every DOCX in a job's Internal/Processing/
folder before PDF conversion. Ensures the proposal package adheres to Format Guide
before it ships.

Actions applied to each DOCX:
  · Table headers → Forest Green #2E7D32 fill, white bold text, Calibri 10pt
  · Body text → Slate #455A64, Calibri 10pt
  · Numbered lists → bullet points (governance: bullets only, no numbered lists)
  · Table rows → cant_split=True (no mid-row page breaks)
  · Paragraph before table → keep_with_next=True (no orphan header)
  · Trailing empty paragraphs + stray page breaks → removed
  · Blue #1976D2, Purple #7B1FA2, Teal #00897B (quote-reserved) → recolored to Slate
  · Pure black text → recolored to Slate

Governance source: 1-Bids/Format Guide.md, memory ### Branding
Never touches: image positioning, logos, footer text, page numbering
                (those are baked into Letterhead.docx)

Usage:
    python "Proposal Packet Rebuilder.py" --folder "path/to/Processing"

Author: GCC LV Div27-28
"""

from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path("/home/user/AI_DRIVE_OUTPUT")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ==========================================================================
# Brand constants
# ==========================================================================

FOREST_GREEN = RGBColor(0x2E, 0x7D, 0x32)
SLATE = RGBColor(0x45, 0x5A, 0x64)
WARM_GOLD = RGBColor(0xD4, 0xAF, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BANNED_COLORS = {
    # Maps banned → replacement
    (0x19, 0x76, 0xD2): SLATE,   # Blue #1976D2 — auto-disqualifies on brand
    (0x7B, 0x1F, 0xA2): SLATE,   # Purple #7B1FA2
    (0x00, 0x89, 0x7B): SLATE,   # Teal #00897B — reserved for quotes
    (0x00, 0x00, 0x00): SLATE,   # Pure black → Slate per governance
}

BODY_FONT = "Calibri"


# ==========================================================================
# Fix helpers
# ==========================================================================

def _set_cell_shading(cell, hex_color):
    """Set cell background color via XML shading element."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing shading
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _is_header_row(row):
    """Detect header row heuristically: first row + any bold text."""
    if len(row.table.rows) > 0 and row is row.table.rows[0]:
        return True
    return False


def fix_tables(doc):
    """Apply brand styling to all tables."""
    for tbl in doc.tables:
        for ri, row in enumerate(tbl.rows):
            is_header = (ri == 0)
            # Prevent row-split page breaks
            trPr = row._tr.get_or_add_trPr()
            for old in trPr.findall(qn("w:cantSplit")):
                trPr.remove(old)
            cant_split = OxmlElement("w:cantSplit")
            trPr.append(cant_split)

            for cell in row.cells:
                if is_header:
                    _set_cell_shading(cell, "2E7D32")
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = BODY_FONT
                            run.font.size = Pt(10)
                            run.font.bold = True
                            run.font.color.rgb = WHITE
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.name is None:
                                run.font.name = BODY_FONT
                            if run.font.size is None:
                                run.font.size = Pt(10)
                            # Recolor banned colors
                            if run.font.color and run.font.color.rgb is not None:
                                rgb_tuple = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                                if rgb_tuple in BANNED_COLORS:
                                    run.font.color.rgb = BANNED_COLORS[rgb_tuple]


def convert_numbered_to_bullets(doc):
    """Convert numbered list items to bullets (per governance: bullets only)."""
    for para in doc.paragraphs:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        # Change numId to a bullet numId. Easiest approach: set to the bullet style
        # by removing numPr and applying the paragraph's ListBullet style if present.
        pPr.remove(numPr)
        try:
            para.style = doc.styles["List Bullet"]
        except KeyError:
            pass  # Fallback: just remove numbering


def fix_body_colors(doc):
    """Recolor any body text using banned colors → Slate."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color and run.font.color.rgb is not None:
                rgb = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                if rgb in BANNED_COLORS:
                    run.font.color.rgb = BANNED_COLORS[rgb]


def keep_with_next_before_tables(doc):
    """Set keep_with_next on paragraphs immediately preceding a table."""
    body = doc.element.body
    for i, child in enumerate(list(body)):
        if child.tag == qn("w:tbl") and i > 0:
            prev = body[i - 1]
            if prev.tag == qn("w:p"):
                pPr = prev.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    prev.insert(0, pPr)
                # Remove any existing keepNext
                for k in pPr.findall(qn("w:keepNext")):
                    pPr.remove(k)
                pPr.append(OxmlElement("w:keepNext"))


def strip_trailing_empty_paragraphs(doc):
    """Remove empty paragraphs + stray page breaks at the end of the document."""
    body = doc.element.body
    # Iterate from the end
    for elem in list(body)[::-1]:
        if elem.tag != qn("w:p"):
            break
        # Check if paragraph has any meaningful text
        text = ''.join([t.text or '' for t in elem.iter(qn("w:t"))])
        if text.strip():
            break
        # Remove empty paragraph (including those with only a page break)
        body.remove(elem)


# ==========================================================================
# Orchestrator
# ==========================================================================

def rebuild_docx(path):
    """Apply all fixes to a single DOCX in place (overwrites)."""
    doc = Document(str(path))
    fix_tables(doc)
    convert_numbered_to_bullets(doc)
    fix_body_colors(doc)
    keep_with_next_before_tables(doc)
    strip_trailing_empty_paragraphs(doc)
    doc.save(str(path))


def rebuild_folder(folder):
    folder = Path(folder)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    docx_files = [p for p in folder.rglob("*.docx") if not p.name.startswith("~$")]
    if not docx_files:
        print(f"[WARN] No DOCX files in {folder}")
        return []

    processed = []
    for docx_path in docx_files:
        try:
            rebuild_docx(docx_path)
            processed.append(docx_path)
            print(f"[OK] Rebuilt: {docx_path.name}")
        except Exception as e:
            print(f"[ERR] {docx_path.name}: {e}")

    print(f"\n[DONE] {len(processed)} of {len(docx_files)} file(s) rebuilt to brand standard")
    return processed


def main():
    p = argparse.ArgumentParser(description="Batch-apply GCC brand formatting to DOCX files")
    p.add_argument("--folder", required=True, help="Path to folder (DOCX files scanned recursively)")
    args = p.parse_args()
    rebuild_folder(args.folder)


if __name__ == "__main__":
    main()
